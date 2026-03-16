"""
Markdown to DOCX Converter
Converts the Remote Backend & WebSocket Guide from .md to .docx format
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Style the hyperlink
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX with formatting"""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Split content into lines
    lines = content.split('\n')
    
    in_code_block = False
    code_block_content = []
    in_table = False
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_block_content)
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                
                code_block_content = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
                code_block_content = []
            i += 1
            continue
        
        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue
        
        # Handle tables
        if line.startswith('|') and '|' in line:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table and not line.startswith('|'):
            # End of table
            create_table(doc, table_lines)
            table_lines = []
            in_table = False
        
        # Skip empty lines
        if not line.strip():
            if i > 0 and lines[i-1].strip():  # Add paragraph break only after content
                doc.add_paragraph()
            i += 1
            continue
        
        # Handle headers
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            text = line.lstrip('#').strip()
            
            p = doc.add_paragraph(text)
            if level == 1:
                p.style = 'Heading 1'
                p.runs[0].font.size = Pt(24)
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            elif level == 2:
                p.style = 'Heading 2'
                p.runs[0].font.size = Pt(18)
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(0, 102, 204)
            elif level == 3:
                p.style = 'Heading 3'
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(51, 102, 153)
            else:
                p.style = 'Heading 4'
                p.runs[0].font.size = Pt(12)
                p.runs[0].font.bold = True
            
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.border_bottom = True
            i += 1
            continue
        
        # Handle lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(text, style='List Bullet')
            format_inline_markdown(p)
            i += 1
            continue
        
        if re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            p = doc.add_paragraph(text, style='List Number')
            format_inline_markdown(p)
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph(line)
        format_inline_markdown(p)
        i += 1
    
    # Handle remaining table if any
    if in_table and table_lines:
        create_table(doc, table_lines)
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Successfully converted to: {docx_file}")

def create_table(doc, table_lines):
    """Create a formatted table from markdown table lines"""
    # Remove empty lines
    table_lines = [line for line in table_lines if line.strip()]
    
    if len(table_lines) < 2:
        return
    
    # Parse table
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.split('|')]
        cells = [c for c in cells if c]  # Remove empty cells
        if cells and not all(c.startswith('-') for c in cells):  # Skip separator line
            rows.append(cells)
    
    if not rows:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Fill table
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            
            # Header row formatting
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                cell._element.get_or_add_tcPr().append(
                    OxmlElement('w:shd')
                )
                cell._element.get_or_add_tcPr().find(qn('w:shd')).set(qn('w:fill'), '4472C4')

def format_inline_markdown(paragraph):
    """Format inline markdown (bold, italic, code)"""
    text = paragraph.text
    paragraph.clear()
    
    # Split by inline code first
    parts = re.split(r'(`[^`]+`)', text)
    
    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            # Inline code
            code_text = part[1:-1]
            run = paragraph.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(199, 37, 78)
        else:
            # Handle bold and italic
            sub_parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', part)
            for sub_part in sub_parts:
                if sub_part.startswith('**') and sub_part.endswith('**'):
                    run = paragraph.add_run(sub_part[2:-2])
                    run.bold = True
                elif sub_part.startswith('*') and sub_part.endswith('*'):
                    run = paragraph.add_run(sub_part[1:-1])
                    run.italic = True
                else:
                    # Check for URLs
                    url_pattern = r'(https?://[^\s]+)'
                    if re.search(url_pattern, sub_part):
                        url_parts = re.split(url_pattern, sub_part)
                        for url_part in url_parts:
                            if re.match(url_pattern, url_part):
                                run = paragraph.add_run(url_part)
                                run.font.color.rgb = RGBColor(0, 0, 255)
                                run.underline = True
                            else:
                                paragraph.add_run(url_part)
                    else:
                        paragraph.add_run(sub_part)

if __name__ == '__main__':
    import os
    
    # File paths
    md_file = os.path.join('docs', 'Remote-Backend-WebSocket-Guide.md')
    docx_file = os.path.join('docs', 'Remote-Backend-WebSocket-Guide.docx')
    
    # Check if markdown file exists
    if not os.path.exists(md_file):
        print(f"❌ Error: {md_file} not found!")
        print("Please make sure the markdown file exists in the docs folder.")
        exit(1)
    
    print("🔄 Converting Markdown to DOCX...")
    print(f"📄 Input: {md_file}")
    print(f"📝 Output: {docx_file}")
    print()
    
    try:
        convert_markdown_to_docx(md_file, docx_file)
        print()
        print("🎉 Conversion complete!")
        print(f"📂 Open: {os.path.abspath(docx_file)}")
    except Exception as e:
        print(f"❌ Error during conversion: {str(e)}")
        import traceback
        traceback.print_exc()

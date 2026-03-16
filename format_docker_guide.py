"""
Format Docker Implementation Guide to DOCX with proper formatting
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def add_formatted_content(doc, md_file):
    """Convert markdown to formatted DOCX"""
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                code_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # H1 - Main headings
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            p.runs[0].font.size = Pt(18)
            p.runs[0].bold = True
        
        # H2 - Section headings
        elif line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            p.runs[0].font.color.rgb = RGBColor(0, 102, 204)
            p.runs[0].font.size = Pt(14)
            p.runs[0].bold = True
        
        # H3 - Subsection headings
        elif line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            p.runs[0].font.color.rgb = RGBColor(51, 51, 51)
            p.runs[0].font.size = Pt(12)
            p.runs[0].bold = True
        
        # H4
        elif line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_heading(text, level=4)
            p.runs[0].font.size = Pt(11)
            p.runs[0].bold = True
        
        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph('_' * 80)
        
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Process inline formatting
            text = process_inline_formatting(text)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            p.paragraph_format.left_indent = Inches(0.25)
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line).strip()
            text = process_inline_formatting(text)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
        
        # Tables (simple detection)
        elif '|' in line and line.strip().startswith('|'):
            # Skip table formatting for simplicity, add as preformatted
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
        
        # Regular paragraph
        elif line.strip():
            text = process_inline_formatting(line)
            p = doc.add_paragraph()
            add_formatted_text(p, text)
            p.paragraph_format.space_after = Pt(6)
        
        # Empty line
        else:
            doc.add_paragraph()
        
        i += 1

def process_inline_formatting(text):
    """Process markdown inline formatting"""
    # Don't process if it's already formatted
    return text

def add_formatted_text(paragraph, text):
    """Add text with inline formatting to paragraph"""
    # Handle **bold**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(199, 37, 78)
        else:
            # Handle inline code `code`
            code_parts = re.split(r'(`[^`]+`)', part)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = paragraph.add_run(cp[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(199, 37, 78)
                elif cp:
                    paragraph.add_run(cp)

def main():
    print("🔄 Creating formatted DOCX...")
    
    md_file = 'docs/DOCKER_IMPLEMENTATION_GUIDE.md'
    docx_file = 'docs/DOCKER_IMPLEMENTATION_GUIDE.docx'
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_heading('Docker Implementation Guide', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    title.runs[0].font.size = Pt(24)
    
    # Add subtitle
    subtitle = doc.add_paragraph('FYP Video Processing Application')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    
    # Add date
    date_p = doc.add_paragraph('December 27, 2025')
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(11)
    date_p.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()
    
    # Add formatted content
    add_formatted_content(doc, md_file)
    
    # Save document
    doc.save(docx_file)
    
    print(f"✅ Formatted DOCX created: {docx_file}")
    print(f"📂 Full path: C:\\Users\\JAD\\Documents\\FYP\\{docx_file}")

if __name__ == '__main__':
    main()
